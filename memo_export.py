from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from native_charts import update_native_pies, parse_label_colors
from fund_info import effective, fmt_money, fmt_ext


PROJECT_BALANCE_TEMPLATE = Path(
    r"C:\Users\jfrischman\OneDrive - GCM Grosvenor\Credit\Credit Secondaries\1. Investments\1. IC Approved & Closed Transactions\Project Balance\Project Balance IC Memo v1.docx"
)

# Ordered asset classes for the Asset Type By Fund table.
_ASSET_CLASS_ORDER = ["Corporate Lending", "ABS", "Special Situations"]

# Security type → asset class (matches the app's dropdown)
_SEC_TO_AC = {
    "Direct Lending": "Corporate Lending",
    "Other Senior Lending": "Corporate Lending",
    "Opportunistic / Junior": "Corporate Lending",
    "Distressed": "Corporate Lending",
    "Corporate Equity": "Corporate Lending",
    "CLOs": "ABS",
    "Regulatory Capital": "ABS",
    "Commercial RE (Debt)": "ABS",
    "Residential RE": "ABS",
    "Consumer": "ABS",
    "Hard Assets": "ABS",
    "Specialty Lending": "ABS",
    "Commercial RE (Equity)": "Special Situations",
    "Commercial RE (Non-Perf)": "Special Situations",
    "Equity": "Special Situations",
}


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def _fmt_pct(value: float) -> str:
    return f"{value * 100:.1f}%"


def _fmt_pct0(value: float) -> str:
    return f"{value * 100:.0f}%"


def _fmt_money(value: float) -> str:
    if value is None:
        return "-"
    try:
        num = float(value)
    except Exception:
        return "-"
    if abs(num) >= 1_000_000:
        return f"${num / 1_000_000:.1f}m"
    return f"${num:,.0f}"


def _set_cell_text(cell, text: str, align: Optional[WD_ALIGN_PARAGRAPH] = None,
                   bold: bool = False, italic: bool = False, font_pt: int = 8):
    cell.text = str(text)
    for p in cell.paragraphs:
        if align is not None:
            p.alignment = align
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(font_pt)
            if bold:
                run.font.bold = True
            if italic:
                run.font.italic = True


def _set_cell_shading(cell, fill_hex: str = "D9D9D9"):
    """Apply a background fill color to a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_no_wrap(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for el in tcPr.findall(qn("w:noWrap")):
        tcPr.remove(el)
    tcPr.append(OxmlElement("w:noWrap"))


def _format_table(table, center_from_col: int, font_pt: int = 8,
                  wide_col: Optional[int] = None, wide_width_in: float = 1.9):
    for row in table.rows:
        cells = row.cells
        for ci, cell in enumerate(cells):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(font_pt)
            if ci >= center_from_col:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if wide_col is not None and ci == wide_col:
                cell.width = Inches(wide_width_in)
                _set_cell_no_wrap(cell)


def _find_project_balance_template(project: Dict[str, Any]) -> Path:
    memo_path = project.get("memo_file_path") or ""
    if memo_path and Path(memo_path).exists():
        return Path(memo_path)
    return PROJECT_BALANCE_TEMPLATE


# ---------------------------------------------------------------------------
# Investment Summary box (table[0]) – label-search so merged cells don't bite
# ---------------------------------------------------------------------------

def _apply_summary_stats(doc: Document, result: Dict[str, Any]) -> None:
    """Update the Investment Summary box using label search (robust to merged cells)."""
    if not doc.tables:
        return
    t = doc.tables[0]
    top = result.get("top_concentration") or {}
    positions = result.get("top_positions") or []
    fund_profiles = result.get("fund_profiles") or []

    # Build a flat map: label_text -> (row_idx, col_idx) for every cell
    label_to_pos: Dict[str, tuple] = {}
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt:
                label_to_pos[txt] = (ri, ci)

    # Apply alignment and Calibri 9 directly via XML (bypasses table-style inheritance).
    # Use row._tr.tc_lst to get actual cells with correct grid column positions.
    # Rule: title rows (0-1) → all LEFT; data rows → even cols LEFT, odd cols CENTER.
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    from lxml import etree as _et

    def _force_align(p_el, jc_val):
        pPr = p_el.find(f"{{{W}}}pPr")
        if pPr is None:
            pPr = _et.SubElement(p_el, f"{{{W}}}pPr")
        for old in pPr.findall(f"{{{W}}}jc"):
            pPr.remove(old)
        jc = _et.SubElement(pPr, f"{{{W}}}jc")
        jc.set(f"{{{W}}}val", jc_val)

    for ri, row in enumerate(t.rows):
        # Use tc index (not grid column) — avoids gridSpan miscounting.
        # Table pattern: i=0,2,4 → labels (LEFT); i=1,3,5 → values (CENTER).
        # Title rows (ri 0-1) and any merged title cell → always LEFT.
        for i, tc in enumerate(row._tr.tc_lst):
            jc_val = "left" if (ri <= 1 or i % 2 == 0) else "center"
            for p_el in tc.findall(f"{{{W}}}p"):
                _force_align(p_el, jc_val)
                for r_el in p_el.findall(f".//{{{W}}}r"):
                    rPr = r_el.find(f"{{{W}}}rPr")
                    if rPr is None:
                        rPr = _et.SubElement(r_el, f"{{{W}}}rPr")
                    for tag in ("rFonts", "sz", "szCs"):
                        for old in rPr.findall(f"{{{W}}}{tag}"):
                            rPr.remove(old)
                    rf = _et.SubElement(rPr, f"{{{W}}}rFonts")
                    rf.set(f"{{{W}}}ascii", "Calibri")
                    rf.set(f"{{{W}}}hAnsi", "Calibri")
                    sz = _et.SubElement(rPr, f"{{{W}}}sz")
                    sz.set(f"{{{W}}}val", "18")  # 9pt = 18 half-points

    # Update specific value cells (these calls also write the new text + CENTER)
    def set_next(label: str, value: str):
        pos = label_to_pos.get(label)
        if pos is None:
            return
        ri, ci = pos
        try:
            _set_cell_text(t.rows[ri].cells[ci + 1], value,
                           align=WD_ALIGN_PARAGRAPH.CENTER, font_pt=9)
        except (IndexError, Exception):
            pass

    set_next("Top 1 Position", _fmt_pct(top.get("top_1", 0)))
    set_next("Top 5 Position", _fmt_pct(top.get("top_5", 0)))
    gt20 = sum(1 for p in positions if float(p.get("value", 0)) > 0.20)
    set_next("Positions >20%", str(gt20))
    set_next("Underlying Funds", str(len(fund_profiles)))
    set_next("Underlying Investments", str(len(positions)))


# ---------------------------------------------------------------------------
# Concentration table (table[2]) – dynamic rows (Total + one per fund)
# ---------------------------------------------------------------------------

def _compute_top_concentration(items: Sequence[Dict[str, Any]],
                               fund_weight: float = 1.0) -> Dict[str, float]:
    """Compute top-N concentration. If fund_weight < 1, values are project-shares;
    divide by fund_weight to get fund-level percentages for the per-fund rows."""
    ordered = sorted(items, key=lambda x: float(x.get("value") or x.get("percentage") or 0), reverse=True)
    scale = (1.0 / fund_weight) if fund_weight > 0 else 1.0
    vals = [float(x.get("value") or x.get("percentage") or 0) * scale for x in ordered]
    return {
        "top_1": sum(vals[:1]),
        "top_3": sum(vals[:3]),
        "top_5": sum(vals[:5]),
        "top_10": sum(vals[:10]),
        "remaining": max(0.0, 1.0 - sum(vals[:10])),
    }


def _rebuild_concentration_table(table, result: Dict[str, Any]) -> None:
    """Completely rebuild the concentration table with a clean 6-column grid.

    The memo template may have a wide gridCol structure (e.g. 29 columns with
    large cell spans) that causes partial overwrites when only 6 values are
    written. We avoid this by wiping the grid and all rows and starting fresh.
    """
    import copy
    from lxml import etree as _et
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    fund_profiles = result.get("fund_profiles") or []
    top = result.get("top_concentration") or {}

    headers = ["Top Positions", "Top 1", "Top 3", "Top 5", "Top 10", "Remaining"]
    data_rows = [["Total", _fmt_pct(top.get("top_1", 0)), _fmt_pct(top.get("top_3", 0)),
                  _fmt_pct(top.get("top_5", 0)), _fmt_pct(top.get("top_10", 0)),
                  _fmt_pct(top.get("remaining", 0))]]
    for fp in fund_profiles:
        w = float(fp.get("weight") or 0)
        p = _compute_top_concentration(fp.get("position_exposure") or [], fund_weight=w)
        data_rows.append([str(fp.get("fund_name") or "Fund"),
                          _fmt_pct(p["top_1"]), _fmt_pct(p["top_3"]),
                          _fmt_pct(p["top_5"]), _fmt_pct(p["top_10"]),
                          _fmt_pct(p["remaining"])])

    tbl = table._tbl
    n_cols = 6
    col_widths = ["2200", "1100", "1100", "1100", "1100", "1200"]  # twips

    # Reset tblGrid to exactly 6 columns
    tblGrid = tbl.find(f"{{{W}}}tblGrid")
    if tblGrid is None:
        tblGrid = _et.SubElement(tbl, f"{{{W}}}tblGrid")
    for gc in list(tblGrid.findall(f"{{{W}}}gridCol")):
        tblGrid.remove(gc)
    for w in col_widths:
        gc = _et.SubElement(tblGrid, f"{{{W}}}gridCol")
        gc.set(f"{{{W}}}w", w)

    # Grab a prototype cell from the existing table for border/shading styles
    existing_trs = list(tbl.findall(f"{{{W}}}tr"))
    proto_tc = None
    for tr in existing_trs:
        tcs = tr.findall(f"{{{W}}}tc")
        for tc in tcs:
            # Remove any gridSpan so the prototype is a plain 1-col cell
            tcPr = tc.find(f"{{{W}}}tcPr")
            if tcPr is not None:
                for gs in tcPr.findall(f"{{{W}}}gridSpan"):
                    tcPr.remove(gs)
            proto_tc = copy.deepcopy(tc)
            # Strip text content
            for p_el in list(proto_tc.findall(f"{{{W}}}p")):
                proto_tc.remove(p_el)
            _et.SubElement(proto_tc, f"{{{W}}}p")
            break
        if proto_tc is not None:
            break

    if proto_tc is None:
        proto_tc = OxmlElement("w:tc")
        proto_tc.append(OxmlElement("w:p"))

    def _make_tr(n: int):
        # OxmlElement("w:tr") produces a CT_Row instance (has tc_lst, used by row.cells)
        tr = OxmlElement("w:tr")
        for _ in range(n):
            tr.append(copy.deepcopy(proto_tc))
        return tr

    # Remove all existing rows and rebuild
    for tr in list(existing_trs):
        tbl.remove(tr)

    # Header row (italic column titles, light green background)
    hdr_tr = _make_tr(n_cols)
    tbl.append(hdr_tr)
    hdr_row = table.rows[0]
    for ci, hdr in enumerate(headers):
        bold = ci == 0
        _set_cell_text(hdr_row.cells[ci], hdr,
                       align=WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else None,
                       bold=bold, italic=not bold, font_pt=8)
        _set_cell_shading(hdr_row.cells[ci], "DDE8CB")  # light green like template
        hdr_row.cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for row_vals in data_rows:
        tbl.append(_make_tr(n_cols))
    for i, row_vals in enumerate(data_rows):
        row = table.rows[i + 1]
        for ci, val in enumerate(row_vals):
            cell = row.cells[ci]
            _set_cell_text(cell, val,
                           align=WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else None,
                           font_pt=8)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Remove any inherited fill from the prototype cell
            tcPr = cell._tc.get_or_add_tcPr()
            for shd in tcPr.findall(qn("w:shd")):
                tcPr.remove(shd)

    _format_table(table, center_from_col=1)


# ---------------------------------------------------------------------------
# Asset Type By Fund table (table[3]) – fully dynamic, ABS support, gray headers
# ---------------------------------------------------------------------------

def _rebuild_asset_type_table(table, result: Dict[str, Any]) -> None:
    """Rebuild the Asset Type By Fund table to match memo format:
    - Column header row (Asset Class | Sub-Asset Class | Total | fund...)
    - LP NAV row
    - Gray + bold asset class header row for each present asset class
    - Italic sub-rows for each present security type under that asset class
    Dynamic: rows added/removed as security types change; fund columns follow fund count.
    """
    import copy
    categories = result.get("categories") or {}
    fund_profiles = result.get("fund_profiles") or []

    asset_vals = {it["label"]: float(it.get("value") or it.get("percentage") or 0)
                  for it in categories.get("asset_class", [])}
    sub_src = categories.get("security_type") or []
    sub_vals = {it["label"]: float(it.get("value") or it.get("percentage") or 0)
                for it in sub_src}

    fund_weights = [float(f.get("weight") or 0) for f in fund_profiles]
    total_w = sum(fund_weights) or 1.0
    fund_weight_pcts = [w / total_w for w in fund_weights]
    n_funds = len(fund_profiles)
    fund_names = [str(fp.get("fund_name") or "Fund") for fp in fund_profiles]

    def fund_ac(fi: int, label: str) -> str:
        cats = (fund_profiles[fi].get("categories") or {}).get("asset_class", [])
        m = {it["label"]: float(it.get("value") or it.get("percentage") or 0) for it in cats}
        return _fmt_pct0(m.get(label, 0.0))

    def fund_sub(fi: int, label: str) -> str:
        fp_cats = fund_profiles[fi].get("categories") or {}
        cats = fp_cats.get("security_type") or []
        m = {it["label"]: float(it.get("value") or it.get("percentage") or 0) for it in cats}
        return _fmt_pct0(m.get(label, 0.0))

    # Row specs: (row_type, values)
    # row_type: "header" | "lpnav" | "ac_header" | "sub"
    specs: List[tuple] = []
    specs.append(("header", ["Asset Class", "Sub-Asset Class", "Total"] + fund_names))
    specs.append(("lpnav",  ["LP NAV", "", "100%"] + [_fmt_pct0(w) for w in fund_weight_pcts]))

    present_acs = {_SEC_TO_AC.get(s, "") for s in sub_vals}
    for ac in _ASSET_CLASS_ORDER:
        if ac not in asset_vals and ac not in present_acs:
            continue
        specs.append(("ac_header", [ac, "", _fmt_pct0(asset_vals.get(ac, 0.0))]
                      + [fund_ac(fi, ac) for fi in range(n_funds)]))
        _sub_order = [
            "Direct Lending", "Other Senior Lending", "Opportunistic / Junior",
            "Distressed", "Corporate Equity",
            "CLOs", "Regulatory Capital", "Commercial RE (Debt)", "Residential RE",
            "Consumer", "Hard Assets", "Specialty Lending",
            "Commercial RE (Equity)", "Commercial RE (Non-Perf)", "Equity",
        ]
        _sub_pos = {s: i for i, s in enumerate(_sub_order)}
        subs = sorted(
            [s for s, v in sub_vals.items() if _SEC_TO_AC.get(s) == ac and v > 0],
            key=lambda s: _sub_pos.get(s, 999)
        )
        for sub in subs:
            specs.append(("sub", [ac, sub, _fmt_pct0(sub_vals.get(sub, 0.0))]
                          + [fund_sub(fi, sub) for fi in range(n_funds)]))

    # --- Sync the table's column grid to match the required fund count ---
    from lxml import etree as _lxml_et
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    tbl = table._tbl
    n_cols_needed = 3 + n_funds   # Asset Class | Sub-Asset Class | Total | fund...

    tblGrid = tbl.find(f"{{{W}}}tblGrid")
    if tblGrid is not None:
        existing = tblGrid.findall(f"{{{W}}}gridCol")
        n_cols_current = len(existing)
        if n_cols_needed != n_cols_current:
            # Derive a sensible width for the fund columns from the last existing one
            fund_w = existing[-1].get(f"{{{W}}}w", "1400") if existing else "1400"
            # Remove all and rebuild with correct count
            for gc in list(existing):
                tblGrid.remove(gc)
            # Approximate widths: issuer col wide, sub-asset wide, then narrower numerics
            col_widths = ["3200", "2600"] + ["1200"] * (1 + n_funds)
            for w in col_widths[:n_cols_needed]:
                gc = _lxml_et.SubElement(tblGrid, f"{{{W}}}gridCol")
                gc.set(f"{{{W}}}w", w)

    # --- Build a prototype cell (copy tcPr from a data cell in the original row) ---
    header_tr = list(tbl.tr_lst)[0]
    orig_tcs = header_tr.findall(f"{{{W}}}tc")
    proto_tc = copy.deepcopy(orig_tcs[2]) if len(orig_tcs) > 2 else copy.deepcopy(orig_tcs[-1])
    # Clear text content from the prototype, keep only tcPr
    for p_el in list(proto_tc.findall(f"{{{W}}}p")):
        proto_tc.remove(p_el)
    empty_p = _lxml_et.SubElement(proto_tc, f"{{{W}}}p")

    def _make_tr(n_cells: int) -> "_lxml_et._Element":
        tr = _lxml_et.Element(f"{{{W}}}tr")
        for _ in range(n_cells):
            tr.append(copy.deepcopy(proto_tc))
        return tr

    # Remove all existing rows, then add fresh rows with the correct cell count
    for tr in list(tbl.tr_lst):
        tbl.remove(tr)
    for _ in specs:
        tbl.append(_make_tr(n_cols_needed))

    # --- Fill values with appropriate formatting per row type ---
    for row_i, (rtype, vals) in enumerate(specs):
        row = table.rows[row_i]
        is_ac_header = (rtype == "ac_header")
        is_sub = (rtype == "sub")
        is_col_header = (rtype == "header")

        for ci, val in enumerate(vals):
            if ci >= len(row.cells):
                break
            align = WD_ALIGN_PARAGRAPH.CENTER if ci >= 2 else None
            bold = is_ac_header or is_col_header
            italic = is_sub
            _set_cell_text(row.cells[ci], val, align=align, bold=bold, italic=italic)
            row.cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if is_ac_header:
                _set_cell_shading(row.cells[ci], "D9D9D9")
        if ci == 1 and ci < len(row.cells):
            row.cells[1].width = Inches(1.9)
            _set_cell_no_wrap(row.cells[1])


def _find_table_after_title(doc: Document, title: str):
    """
    Return the data table that follows a section title, scanning body elements
    in document order. Handles two common memo structures:
      - Title in a standalone paragraph → return the next table.
      - Title in the first row of a small header table → return the table after that.
    Returns None if the title is not found.
    """
    from docx.table import Table as _Table
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    title_lower = title.lower()
    found = False
    for el in doc.element.body:
        is_tbl = el.tag == f"{{{W}}}tbl"
        if is_tbl:
            if found:
                return _Table(el, doc)
            # Check whether the table's first row is itself a title header
            rows = [c for c in el if c.tag == f"{{{W}}}tr"]
            if rows:
                first_row_text = "".join(
                    t.text or "" for t in rows[0].iter() if t.tag == f"{{{W}}}t"
                ).lower()
                if title_lower in first_row_text:
                    found = True   # next table is the data table
        else:
            text = "".join(t.text or "" for t in el.iter() if t.tag == f"{{{W}}}t").lower()
            if title_lower in text:
                found = True
    return None


def _remove_duplicate_asset_tables(doc: Document, keep_el) -> None:
    """Delete any tables other than keep_el whose first row contains both
    'asset class' and 'sub-asset' cell text — i.e. stale Asset Type By Fund
    tables left behind by previous update runs."""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    body = doc.element.body
    to_remove = []
    for el in body:
        if el.tag != f"{{{W}}}tbl" or el is keep_el:
            continue
        rows = [c for c in el if c.tag == f"{{{W}}}tr"]
        if not rows:
            continue
        cell_texts = [
            "".join(t.text or "" for t in tc.iter() if t.tag == f"{{{W}}}t").lower()
            for tc in rows[0] if tc.tag == f"{{{W}}}tc"
        ]
        if (any("asset class" in s for s in cell_texts) and
                any("sub-asset" in s for s in cell_texts)):
            to_remove.append(el)
    for el in to_remove:
        body.remove(el)


def _find_concentration_table(doc: Document):
    """Find the Top Positions concentration table by title search, falling back to index 2."""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for el in doc.element.body:
        if el.tag != f"{{{W}}}tbl":
            continue
        rows = [c for c in el if c.tag == f"{{{W}}}tr"]
        if not rows:
            continue
        first_text = "".join(
            t.text or "" for t in rows[0].iter() if t.tag == f"{{{W}}}t"
        ).lower()
        if "top position" in first_text or "top 1" in first_text:
            from docx.table import Table as _Table
            return _Table(el, doc)
    return doc.tables[2] if len(doc.tables) >= 3 else None


def _apply_exposure_tables(doc: Document, result: Dict[str, Any]) -> None:
    _apply_summary_stats(doc, result)
    conc_tbl = _find_concentration_table(doc)
    if conc_tbl is not None:
        _rebuild_concentration_table(conc_tbl, result)
    asset_type_tbl = _find_table_after_title(doc, "Asset Type By Fund")
    if asset_type_tbl is None and len(doc.tables) >= 4:
        asset_type_tbl = doc.tables[3]   # positional fallback for Project Balance layout
    if asset_type_tbl is not None:
        _rebuild_asset_type_table(asset_type_tbl, result)
        _remove_duplicate_asset_tables(doc, asset_type_tbl._tbl)


# ---------------------------------------------------------------------------
# Section registry
# ---------------------------------------------------------------------------

def _update_current_portfolio_names(doc: Document, result: Dict[str, Any],
                                    project_name: str = "Project") -> None:
    """Replace the bold company-name line in each Current Portfolio numbered entry
    with the live exposure format: 'Name (X.X% ProjectName, X.X% Fund1, ...)'.
    Finds entries by ilvl=0 List Paragraph paragraphs after the section heading.
    Works even when content is inside tracked-change elements (<w:ins>/<w:moveTo>).
    Updates the top 8 positions; skips paragraphs with no text content."""
    from lxml import etree as _et
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    XML_SPACE = "http://www.w3.org/XML/1998/namespace"

    fund_profiles = result.get("fund_profiles") or []
    top_positions = result.get("top_positions") or []

    # Per-fund position lookup: {(fund_idx, key): fund_level_%}
    fund_pos: Dict[tuple, float] = {}
    for fi, fp in enumerate(fund_profiles):
        w = float(fp.get("weight") or 0)
        for pos in (fp.get("position_exposure") or []):
            key = str(pos.get("key") or pos.get("label") or "").lower()
            val = float(pos.get("value") or 0)
            fund_pos[(fi, key)] = (val / w) if w > 0 else 0.0

    def _fmt(pct: float) -> str:
        return f"{pct * 100:.1f}%"

    def _build_label(pos_item: Dict[str, Any]) -> str:
        name = pos_item.get("label") or ""
        key = str(pos_item.get("key") or name).lower()
        proj_pct = float(pos_item.get("value") or 0)
        parts = [f"{_fmt(proj_pct)} {project_name}"]
        for fi, fp in enumerate(fund_profiles):
            fp_pct = fund_pos.get((fi, key), 0.0)
            if fp_pct > 0:
                abbrev = str(fp.get("abbrev_name") or fp.get("fund_name") or f"Fund {fi+1}")
                parts.append(f"{_fmt(fp_pct)} {abbrev}")
        return f"{name} ({', '.join(parts)})"

    def _all_text(p_el) -> str:
        return "".join(t.text or "" for t in p_el.iter() if t.tag == f"{{{W}}}t")

    def _replace_content(p_el, new_text: str):
        """Keep pPr + bookmarks, replace everything else with a single bold run."""
        keep_tags = {f"{{{W}}}pPr", f"{{{W}}}bookmarkStart", f"{{{W}}}bookmarkEnd"}
        for child in list(p_el):
            if child.tag not in keep_tags:
                p_el.remove(child)
        r = _et.SubElement(p_el, f"{{{W}}}r")
        rPr = _et.SubElement(r, f"{{{W}}}rPr")
        _et.SubElement(rPr, f"{{{W}}}b")
        rFonts = _et.SubElement(rPr, f"{{{W}}}rFonts")
        rFonts.set(f"{{{W}}}ascii", "Calibri")
        rFonts.set(f"{{{W}}}hAnsi", "Calibri")
        sz = _et.SubElement(rPr, f"{{{W}}}sz"); sz.set(f"{{{W}}}val", "18")
        t_el = _et.SubElement(r, f"{{{W}}}t")
        t_el.text = new_text
        t_el.set(f"{{{XML_SPACE}}}space", "preserve")

    # Find Current Portfolio section
    in_section = False
    label_idx = 0
    for para in doc.paragraphs:
        if not in_section:
            if "Current Portfolio" in para.text:
                in_section = True
            continue

        if label_idx >= min(8, len(top_positions)):
            break

        # Check for ilvl=0 (company entry paragraph)
        pPr = para._p.find(f"{{{W}}}pPr")
        if pPr is None:
            continue
        numPr = pPr.find(f"{{{W}}}numPr")
        if numPr is None:
            continue
        ilvl_el = numPr.find(f"{{{W}}}ilvl")
        if ilvl_el is None or ilvl_el.get(f"{{{W}}}val") != "0":
            continue

        # Skip truly empty paragraphs
        if not _all_text(para._p).strip():
            continue

        _replace_content(para._p, _build_label(top_positions[label_idx]))
        label_idx += 1


def _update_fund_summary_table(doc: Document, result: Dict[str, Any]) -> None:
    """Update the fund summary table (table[1]) with live fund information.
    Columns: Fund | LP NAV | Unf. | Commits | Invest End | Term End |
             Extensions | IRR | TVPI | RVPI | DPI | Leverage
    Rows are added/removed to match fund count. Values come from fund_infos
    stored in result['_fund_infos'] (effective value = override ?? parsed)."""
    if len(doc.tables) < 2:
        return
    import copy
    from lxml import etree as _et
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    fund_profiles = result.get("fund_profiles") or []
    fund_infos: Dict[str, Any] = result.get("_fund_infos") or {}
    n_funds = len(fund_profiles)
    table = doc.tables[1]
    tbl = table._tbl

    # Sync row count (header = row 0, data rows = 1..n)
    header_tr = list(tbl.tr_lst)[0]
    existing = list(tbl.tr_lst)[1:]
    for tr in existing[n_funds:]:
        tbl.remove(tr)
    for _ in range(n_funds - len(existing)):
        tbl.append(copy.deepcopy(header_tr))

    # Build and fill one row per fund
    any_perp = False
    perp_notes: List[str] = []

    for fi, fp in enumerate(fund_profiles):
        fname = fp.get("fund_name") or f"Fund {fi + 1}"
        finfo = fund_infos.get(fname) or {}
        fields = finfo.get("fields") or {}
        scale = float(finfo.get("scale_pct") or 100.0)

        def eff(k):
            v = effective(fields.get(k))
            return str(v) if v is not None else ""

        ext_field = fields.get("extensions") or {}
        if ext_field.get("perpetuity"):
            any_perp = True
            note = ext_field.get("perpetuity_note") or ""
            if note and note not in perp_notes:
                perp_notes.append(note)

        vals = [
            fname,
            fmt_money(fields.get("lp_nav"),   scale),
            fmt_money(fields.get("unfunded"),  scale),
            fmt_money(fields.get("commits"),   scale),
            eff("invest_end"),
            eff("term_end"),
            fmt_ext(ext_field),
            eff("irr"),
            eff("tvpi"),
            eff("rvpi"),
            eff("dpi"),
            eff("leverage"),
        ]

        row = table.rows[fi + 1]
        for ci, val in enumerate(vals):
            if ci >= len(row.cells):
                break
            align = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            _set_cell_text(row.cells[ci], val, align=align, font_pt=9)
            row.cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _format_table(table, center_from_col=1, font_pt=9)

    # Add / update perpetuity footnote paragraph immediately after table
    footnote_text = ""
    if any_perp:
        note = perp_notes[0] if perp_notes else "after stated extensions, LPAC may approve additional extensions"
        footnote_text = f"* {note}"

    # Find the paragraph right after table[1] in the body XML
    body = doc.element.body
    tbl_el = table._tbl
    tbl_idx = list(body).index(tbl_el)
    # Check if next element is already a footnote paragraph (starts with *)
    next_el = body[tbl_idx + 1] if tbl_idx + 1 < len(body) else None
    next_text = ""
    if next_el is not None and next_el.tag == f"{{{W}}}p":
        next_text = "".join(t.text or "" for t in next_el.iter() if t.tag == f"{{{W}}}t")

    if footnote_text:
        if next_text.startswith("*"):
            # Update existing footnote
            for t in next_el.findall(f".//{{{W}}}t"):
                t.text = footnote_text
        else:
            # Insert new footnote paragraph
            from lxml import etree as _et2
            fn_p = _et2.Element(f"{{{W}}}p")
            fn_pPr = _et2.SubElement(fn_p, f"{{{W}}}pPr")
            fn_rPr = _et2.SubElement(fn_p, f"{{{W}}}r")
            fn_rPrInner = _et2.SubElement(fn_rPr, f"{{{W}}}rPr")
            fn_sz = _et2.SubElement(fn_rPrInner, f"{{{W}}}sz"); fn_sz.set(f"{{{W}}}val", "10")  # 5pt = 10 half-pts
            fn_szCs = _et2.SubElement(fn_rPrInner, f"{{{W}}}szCs"); fn_szCs.set(f"{{{W}}}val", "10")
            fn_font = _et2.SubElement(fn_rPrInner, f"{{{W}}}rFonts")
            fn_font.set(f"{{{W}}}ascii", "Calibri"); fn_font.set(f"{{{W}}}hAnsi", "Calibri")
            fn_t = _et2.SubElement(fn_rPr, f"{{{W}}}t")
            fn_t.text = footnote_text
            tbl_el.addnext(fn_p)
    elif next_text.startswith("*"):
        # Remove stale footnote (no longer perpetuity)
        body.remove(next_el)


def _apply_portfolio_names(doc: Document, result: Dict[str, Any]) -> None:
    """Wrapper that reads project_name from result['_project_name'] (injected before call)."""
    _update_current_portfolio_names(doc, result, result.get("_project_name") or "Project")


def _update_header_net_return(memo_path: Path, irr: float, moic: float) -> None:
    """Replace the value after 'Base Case Expected Net Return: ' in header1.xml.
    Uses direct ZIP manipulation because the header contains tracked-change runs."""
    import zipfile, shutil, tempfile
    from lxml import etree

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    new_text = f"{irr * 100:.1f}% / {moic:.2f}x"

    tmp = Path(tempfile.mktemp(suffix=".docx"))
    shutil.copy2(memo_path, str(tmp))
    try:
        with zipfile.ZipFile(str(tmp), "r") as zin:
            names = zin.namelist()
            if "word/header1.xml" not in names:
                return
            hdr_bytes = zin.read("word/header1.xml")
            other = {n: zin.read(n) for n in names if n != "word/header1.xml"}

        root = etree.fromstring(hdr_bytes)
        label_text = "Base Case Expected Net Return: "

        # Find the <w:r> whose <w:t> contains the label
        label_run = None
        label_para = None
        for p_el in root.iter(f"{{{W}}}p"):
            for r_el in p_el:
                if r_el.tag != f"{{{W}}}r":
                    continue
                t_el = r_el.find(f"{{{W}}}t")
                if t_el is not None and (t_el.text or "").strip().endswith("Net Return:") or \
                   (t_el is not None and label_text.strip() in (t_el.text or "")):
                    label_run = r_el
                    label_para = p_el
                    break
            if label_run is not None:
                break

        if label_para is None or label_run is None:
            return

        # Copy the rPr from the label run to use as the template for the new value run
        label_rPr = label_run.find(f"{{{W}}}rPr")

        # Remove everything after the label run in this paragraph
        children = list(label_para)
        label_idx = children.index(label_run)
        for child in children[label_idx + 1:]:
            label_para.remove(child)

        # Add a clean single run with the new value
        new_r = etree.SubElement(label_para, f"{{{W}}}r")
        if label_rPr is not None:
            import copy
            new_r.insert(0, copy.deepcopy(label_rPr))
        new_t = etree.SubElement(new_r, f"{{{W}}}t")
        new_t.text = new_text
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

        new_hdr_bytes = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

        with zipfile.ZipFile(str(memo_path), "w", zipfile.ZIP_DEFLATED) as zout:
            zout.writestr("word/header1.xml", new_hdr_bytes)
            for n, data in other.items():
                zout.writestr(n, data)
    finally:
        try:
            tmp.unlink(missing_ok=True)
        except Exception:
            pass


def _render_cashflow_table_image(cf_data: dict) -> bytes:
    """Render the Combined Cashflows table as a PNG (3.67 inches wide)."""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
        import io as _io
    except ImportError:
        raise RuntimeError("matplotlib is required. Run: pip install matplotlib")

    header = cf_data.get("header", {})
    cf_rows = cf_data.get("rows", [])

    # Colors matching typical Excel IC memo table styling
    C_BASE  = "#FFC000"
    C_BEAR  = "#2F75B6"
    C_MGR   = "#1F4E79"
    C_DUR   = "#E2EFDA"
    C_TITLE = "#404040"
    C_ALT   = "#F2F2F2"
    C_GRAY  = "#D9D9D9"

    def _pct(v):  return f"{float(v)*100:.1f}%" if v is not None else ""
    def _x(v):    return f"{float(v):.2f}x"     if v is not None else ""
    def _yrs(v):  return f"{float(v):.2f} Yrs"  if v is not None else ""
    def _cf(v):   return f"{float(v):.2f}"       if v is not None else ""

    # Build the ordered rows (type, label, [6 values])
    data_rows = [
        ("irr",  "Net IRR",       [_pct(header.get(k)) for k in
            ("base_irr_sec","base_irr_fund","bear_irr_sec","bear_irr_fund","mgr_irr_sec","mgr_irr_fund")]),
        ("moic", "Net MOIC",      [_x(header.get(k))   for k in
            ("base_moic_sec","base_moic_fund","bear_moic_sec","bear_moic_fund","mgr_moic_sec","mgr_moic_fund")]),
        ("blank","",              ["", "", "", "", "", ""]),
        ("dur",  "Duration (yrs)",[_yrs(header.get(k)) for k in
            ("base_dur_sec","base_dur_fund","bear_dur_sec","bear_dur_fund","mgr_dur_sec","mgr_dur_fund")]),
    ]
    for row in cf_rows:
        data_rows.append(("data", row["date"], [
            _cf(row.get("base_s")), _cf(row.get("base_f")),
            _cf(row.get("bear_s")), _cf(row.get("bear_f")),
            _cf(row.get("mgr_s")),  _cf(row.get("mgr_f")),
        ]))

    # Row heights in inches
    def rh(rtype):
        return 0.065 if rtype == "blank" else 0.125

    row_heights = [rh(rt) for rt, _, _ in data_rows]

    # Figure dimensions (fixed 3.67" wide)
    FIG_W   = 3.67
    TITLE_H = 0.185
    CASE_H  = 0.150
    SCOL_H  = 0.135
    fig_h = TITLE_H + CASE_H + SCOL_H + sum(row_heights)

    # Column layout: date col + 6 value cols (2 per case)
    COL_W = [0.79, 0.48, 0.48, 0.48, 0.48, 0.48, 0.48]
    COL_X = [sum(COL_W[:i]) for i in range(7)]

    DPI = 200
    fig = plt.figure(figsize=(FIG_W, fig_h), dpi=DPI)
    ax  = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, FIG_W)
    ax.set_ylim(0, fig_h)
    ax.set_aspect("auto")
    ax.axis("off")

    def cell(x, y, w, h, bg, txt="", fs=5.5, fg="black", bold=False, ha="center"):
        ax.add_patch(mpatches.Rectangle(
            (x, y), w, h, facecolor=bg, edgecolor="white",
            linewidth=0.4, transform=ax.transData, clip_on=False))
        if txt:
            tx = x + w * 0.05 if ha == "left" else x + w / 2
            ax.text(tx, y + h / 2, txt, fontsize=fs, ha=ha, va="center",
                    color=fg, fontweight="bold" if bold else "normal",
                    fontfamily="DejaVu Sans", transform=ax.transData, clip_on=True)

    y = fig_h

    # Title row
    y -= TITLE_H
    cell(0, y, FIG_W, TITLE_H, C_TITLE, "Combined Cashflows", fs=7, fg="white", bold=True)

    # Case header row (spans 2 value cols per case)
    y -= CASE_H
    cell(0, y, COL_W[0], CASE_H, C_GRAY)
    cell(COL_X[1], y, COL_W[1] + COL_W[2], CASE_H, C_BASE, "BASE CASE",    fs=6, fg="white", bold=True)
    cell(COL_X[3], y, COL_W[3] + COL_W[4], CASE_H, C_BEAR, "BEAR CASE",    fs=6, fg="white", bold=True)
    cell(COL_X[5], y, COL_W[5] + COL_W[6], CASE_H, C_MGR,  "MANAGER CASE", fs=6, fg="white", bold=True)

    # S CFs / F CFs sub-headers
    y -= SCOL_H
    cell(0, y, COL_W[0], SCOL_H, C_GRAY)
    for ci, (cc, lbl) in enumerate(zip(
        [C_BASE, C_BASE, C_BEAR, C_BEAR, C_MGR, C_MGR],
        ["S CFs", "F CFs", "S CFs", "F CFs", "S CFs", "F CFs"]
    )):
        cell(COL_X[ci + 1], y, COL_W[ci + 1], SCOL_H, cc, lbl, fs=5.5, fg="white", bold=True)

    # Data rows
    for ri, (rtype, label, vals) in enumerate(data_rows):
        rrow_h = row_heights[ri]
        y -= rrow_h
        if rtype == "blank":
            cell(0, y, FIG_W, rrow_h, "white")
            continue
        bg = C_DUR if rtype == "dur" else ("white" if ri % 2 == 0 else C_ALT)
        bold_lbl = rtype in ("irr", "moic", "dur")
        cell(0, y, COL_W[0], rrow_h, bg, label, fs=5.5, bold=bold_lbl, ha="left")
        for ci, val in enumerate(vals):
            cell(COL_X[ci + 1], y, COL_W[ci + 1], rrow_h, bg, val, fs=5.5)

    buf = _io.BytesIO()
    # Embed a title so subsequent runs can re-identify this image via ASCII search
    plt.savefig(buf, format="png", dpi=DPI, metadata={"Title": "Combined Cashflows"})
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _replace_cashflow_image(memo_path: Path, cf_data: dict) -> None:
    """Find the Combined Cashflows EMF image in the docx and replace it with a freshly rendered PNG."""
    import struct, re as _re, zipfile, shutil, tempfile
    from lxml import etree as _et

    if not cf_data.get("rows"):
        return

    # Read all zip entries up front (needed by all search paths)
    with zipfile.ZipFile(str(memo_path), "r") as zin:
        all_files = {n: zin.read(n) for n in zin.namelist()}

    rels_content = all_files.get("word/_rels/document.xml.rels", b"").decode("utf-8")
    rels_root = _et.fromstring(rels_content.encode("utf-8"))

    WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    A  = "http://schemas.openxmlformats.org/drawingml/2006/main"
    R  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    emf_name: Optional[str] = None
    rid: Optional[str] = None

    # Stage 1 & 2: search media binary for text markers
    #   UTF-16-LE = original EMF;  ASCII = PNG we rendered (has Title metadata)
    needle_utf16 = "Combined Cashflows".encode("utf-16-le")
    needle_ascii = b"Combined Cashflows"
    for name, data in all_files.items():
        if name.startswith("word/media/") and (needle_utf16 in data or needle_ascii in data):
            emf_name = name
            break

    if emf_name:
        short_target  = emf_name.replace("word/", "")
        filename_only = emf_name.split("/")[-1]
        for rel in rels_root:
            t = rel.get("Target", "")
            if t == short_target or t.endswith("/" + filename_only):
                rid = rel.get("Id")
                break

    # Stage 3: dimension-based fallback — find floating anchor with cx ≈ 3.67"
    # Handles manually-inserted images and pre-marker PNGs from earlier runs.
    if not rid:
        TARGET_CX    = 3355258  # 3.67 inches in EMU
        CX_TOLERANCE = 65000   # ±0.071 inch
        doc_root_s = _et.fromstring(all_files.get("word/document.xml", b""))
        best_diff: int = CX_TOLERANCE + 1
        best_rid:  Optional[str] = None
        for anchor in doc_root_s.iter(f"{{{WP}}}anchor"):
            ext_el = anchor.find(f"{{{WP}}}extent")
            if ext_el is None:
                continue
            try:
                cx = int(ext_el.get("cx", "0"))
            except (ValueError, TypeError):
                continue
            diff = abs(cx - TARGET_CX)
            if diff < best_diff:
                for b in anchor.findall(f".//{{{A}}}blip"):
                    r = b.get(f"{{{R}}}embed")
                    if r:
                        best_rid  = r
                        best_diff = diff
                        break
        if best_rid:
            rid = best_rid
            for rel in rels_root:
                if rel.get("Id") == rid:
                    t = rel.get("Target", "")
                    if t:
                        emf_name = "word/" + t
                    break

    if not rid or not emf_name:
        raise ValueError(
            "Could not find 'Combined Cashflows' image in the memo. "
            "Make sure the memo contains the original cashflow image or a previous replacement."
        )

    # Derive short_target from whichever search path found emf_name
    short_target = emf_name.replace("word/", "")   # e.g. "media/image2.emf"

    # Render the new PNG
    png_bytes = _render_cashflow_table_image(cf_data)

    # Read PNG dimensions from IHDR (bytes 16-24: width, height as big-endian uint32)
    w_px = struct.unpack(">I", png_bytes[16:20])[0]
    h_px = struct.unpack(">I", png_bytes[20:24])[0]
    DPI = 200
    new_cx = int((w_px / DPI) * 914400)
    new_cy = int((h_px / DPI) * 914400)

    # Assign a new filename for the PNG (avoid reusing the old .emf name)
    existing_nums = [int(m.group(1)) for n in all_files
                     for m in [_re.search(r"word/media/image(\d+)\.", n)] if m]
    new_num  = max(existing_nums) + 1 if existing_nums else 1
    new_png_name   = f"word/media/image{new_num}.png"
    new_png_short  = f"media/image{new_num}.png"

    # Update document.xml — find the anchor that embeds this rId and update extent values
    doc_xml = all_files.get("word/document.xml", b"")
    doc_root = _et.fromstring(doc_xml)
    updated_anchor = False
    for anchor in doc_root.iter(f"{{{WP}}}anchor"):
        blips = anchor.findall(f".//{{{A}}}blip")
        if not any(b.get(f"{{{R}}}embed") == rid for b in blips):
            continue
        # Update blip embed to point to new rId (same rId, target changes in rels)
        # Update <wp:extent cx cy>
        extent = anchor.find(f"{{{WP}}}extent")
        if extent is not None:
            extent.set("cx", str(new_cx))
            extent.set("cy", str(new_cy))
        # Update <a:ext cx cy> inside <a:xfrm>
        for ext in anchor.findall(f".//{{{A}}}ext"):
            parent = ext.getparent()
            if parent is not None and parent.tag == f"{{{A}}}xfrm":
                ext.set("cx", str(new_cx))
                ext.set("cy", str(new_cy))
        updated_anchor = True
        break

    new_doc_xml = _et.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    # Update the relationships file: retarget rid to new PNG
    new_rels = rels_content.replace(
        f'Target="{short_target}"',
        f'Target="{new_png_short}"',
    )

    # Ensure PNG content type is registered in [Content_Types].xml
    ct_xml = all_files.get("[Content_Types].xml", b"").decode("utf-8")
    if 'Extension="png"' not in ct_xml and "extension=\"png\"" not in ct_xml.lower():
        ct_xml = ct_xml.replace(
            "</Types>",
            '<Default Extension="png" ContentType="image/png"/></Types>',
        )

    # Rebuild the zip (drop old EMF, add new PNG, updated rels + doc + content types)
    new_files = {}
    for name, data in all_files.items():
        if name == emf_name:
            continue  # drop old image
        elif name == "word/_rels/document.xml.rels":
            new_files[name] = new_rels.encode("utf-8")
        elif name == "word/document.xml" and updated_anchor:
            new_files[name] = new_doc_xml
        elif name == "[Content_Types].xml":
            new_files[name] = ct_xml.encode("utf-8")
        else:
            new_files[name] = data
    new_files[new_png_name] = png_bytes

    with zipfile.ZipFile(str(memo_path), "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in new_files.items():
            zout.writestr(name, data)


def _apply_model_outputs(doc: Document, result: Dict[str, Any]) -> None:
    """Update Investment Summary table (table[0]) with model outputs."""
    model = result.get("_model_outputs")
    if not model or not doc.tables:
        return
    t = doc.tables[0]
    label_to_pos: Dict[str, tuple] = {}
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt:
                label_to_pos[txt] = (ri, ci)

    def set_next(label: str, value: str):
        pos = label_to_pos.get(label)
        if pos is None:
            return
        ri, ci = pos
        try:
            _set_cell_text(t.rows[ri].cells[ci + 1], value,
                           align=WD_ALIGN_PARAGRAPH.CENTER, font_pt=9)
        except (IndexError, Exception):
            pass

    def fmt_bid(v) -> str:
        return f"{float(v) * 100:.1f}c"

    def fmt_pct(v) -> str:
        return f"{float(v) * 100:.1f}%"

    def fmt_moic(v) -> str:
        return f"{float(v):.2f}x"

    if model.get("gross_bid") is not None:
        set_next("Bid Price", fmt_bid(model["gross_bid"]))
    if model.get("eff_bid") is not None:
        set_next("Effective Price", fmt_bid(model["eff_bid"]))

    base_irr = model.get("base_irr")
    base_moic = model.get("base_moic")
    if base_irr is not None and base_moic is not None:
        set_next("Base IRR / MOIC", f"{fmt_pct(base_irr)} / {fmt_moic(base_moic)}")
    elif base_irr is not None:
        set_next("Base IRR / MOIC", fmt_pct(base_irr))

    mgr_irr = model.get("mgr_irr")
    bear_irr = model.get("bear_irr")
    if mgr_irr is not None and bear_irr is not None:
        set_next("Upside / Downside", f"{fmt_pct(mgr_irr)} / {fmt_pct(bear_irr)}")
    elif mgr_irr is not None:
        set_next("Upside / Downside", fmt_pct(mgr_irr))


SECTION_UPDATERS = {
    "exposures": (_apply_exposure_tables, True),
    "portfolio_names": (_apply_portfolio_names, False),
    "fund_info": (_update_fund_summary_table, False),
    "model_outputs": (_apply_model_outputs, False),
}


def build_memo_export(project: Dict[str, Any], result: Dict[str, Any], output_path: Path) -> Path:
    template = _find_project_balance_template(project)
    doc = Document(str(template))
    result["_project_name"] = project.get("project_name") or "Project"
    result["_fund_infos"] = project.get("fund_infos") or {}
    _apply_exposure_tables(doc, result)
    _apply_portfolio_names(doc, result)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    lc = parse_label_colors(project.get("label_colors") or "")
    update_native_pies(output_path, result.get("categories", {}) or {}, label_colors=lc)
    return output_path


def update_sections_in_file(memo_path, result: Dict[str, Any], sections=("exposures",),
                            project: Dict[str, Any] = None) -> Path:
    memo_path = Path(memo_path)
    if not memo_path.exists():
        raise FileNotFoundError(f"Memo file not found: {memo_path}")
    doc = Document(str(memo_path))
    result["_project_name"] = (project or {}).get("project_name") or "Project"
    result["_fund_infos"] = (project or {}).get("fund_infos") or {}
    refresh_pies = False
    for name in sections:
        updater = SECTION_UPDATERS.get(name)
        if not updater:
            continue
        fn, also_pies = updater
        fn(doc, result)
        refresh_pies = refresh_pies or also_pies
    doc.save(str(memo_path))
    if refresh_pies:
        lc = parse_label_colors((project or {}).get("label_colors") or "")
        update_native_pies(memo_path, result.get("categories", {}) or {}, label_colors=lc)
    # Post-save ZIP-based updates for model outputs
    if "model_outputs" in sections:
        model = result.get("_model_outputs") or {}
        base_irr = model.get("base_irr")
        base_moic = model.get("base_moic")
        if base_irr is not None and base_moic is not None:
            _update_header_net_return(memo_path, base_irr, base_moic)
        cf_data = result.get("_cashflows")
        if cf_data and cf_data.get("rows"):
            _replace_cashflow_image(memo_path, cf_data)
    return memo_path
